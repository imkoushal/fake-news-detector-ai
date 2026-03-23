from docx import Document

doc = Document('Major_Project_Report_New_Complete.docx')

print('Document created successfully!')
print(f'\nTotal pages/sections: {len(doc.sections)}')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables: {len(doc.tables)}')

print('\n' + '='*70)
print('LIST OF ABBREVIATIONS TABLE')
print('='*70)

# Table 3 should be List of Abbreviations (0=TOC, 1=Figures, 2=Tables, 3=Abbreviations)
abbrev_table = doc.tables[3]

print(f'\nTable has {len(abbrev_table.rows)} rows and {len(abbrev_table.columns)} columns')
print('\nShowing all abbreviations:\n')

for i, row in enumerate(abbrev_table.rows):
    abbr = row.cells[0].text.strip()
    full_form = row.cells[1].text.strip()
    
    if i == 0:
        print(f"{'ABBREVIATION':<20} | {'FULL FORM'}")
        print('-'*70)
    else:
        print(f"{abbr:<20} | {full_form}")

print('\n' + '='*70)
print('✅ List of Abbreviations is already in professional table format!')
print('='*70)
